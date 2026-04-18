# Django Backend Implementation for KnowHow AI Document Generation
# Place these files in your Django project

# ============================================================================
# 1. models.py - Document Model
# ============================================================================

from django.db import models
from django.contrib.auth import get_user_model

User = get_user_model()

class Document(models.Model):
    """Document model for storing generated legal documents"""

    FORMAT_CHOICES = [
        ('pdf', 'PDF'),
        ('docx', 'DOCX'),
    ]

    id = models.UUIDField(primary_key=True, default=uuid.uuid4, editable=False)
    user = models.ForeignKey(User, on_delete=models.CASCADE, related_name='documents')
    session = models.ForeignKey('ChatSession', on_delete=models.SET_NULL, null=True, blank=True)

    title = models.CharField(max_length=255)
    content = models.TextField()
    format = models.CharField(max_length=10, choices=FORMAT_CHOICES, default='pdf')

    file = models.FileField(upload_to='documents/%Y/%m/', null=True, blank=True)
    file_size = models.IntegerField(null=True, blank=True)

    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)

    class Meta:
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['-created_at']),
            models.Index(fields=['user', '-created_at']),
        ]

    def __str__(self):
        return f"{self.title} ({self.format})"

    @property
    def file_url(self):
        """Return URL for downloading the document"""
        return f"/api/documents/{self.id}/download/"


# ============================================================================
# 2. serializers.py - API Serializers
# ============================================================================

from rest_framework import serializers
from .models import Document

class DocumentSerializer(serializers.ModelSerializer):
    """Serializer for Document model"""

    file_url = serializers.SerializerMethodField()

    class Meta:
        model = Document
        fields = [
            'id', 'title', 'content', 'format', 'session_id',
            'file_url', 'file_size', 'created_at', 'updated_at'
        ]
        read_only_fields = ['id', 'created_at', 'updated_at', 'file_url', 'file_size']

    def get_file_url(self, obj):
        return obj.file_url


class GenerateDocumentSerializer(serializers.Serializer):
    """Serializer for document generation request"""

    session_id = serializers.UUIDField(required=False, allow_null=True)
    title = serializers.CharField(max_length=255)
    content = serializers.CharField()
    format = serializers.ChoiceField(choices=['pdf', 'docx'], default='pdf')


# ============================================================================
# 3. document_generator.py - PDF/DOCX Generation
# ============================================================================

from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

class DocumentGenerator:
    """Generate PDF and DOCX documents"""

    @staticmethod
    def generate_pdf(title, content):
        """Generate PDF document"""
        buffer = io.BytesIO()

        # Create PDF
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                rightMargin=72, leftMargin=72,
                                topMargin=72, bottomMargin=18)

        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor='#003A80',  # KnowHow blue
            spaceAfter=30,
            alignment=1  # Center
        )

        content_style = ParagraphStyle(
            'CustomContent',
            parent=styles['BodyText'],
            fontSize=12,
            leading=18,
            alignment=4  # Justify
        )

        # Build document
        story = []

        # Add title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.5 * inch))

        # Add content (split by paragraphs)
        for paragraph in content.split('\n'):
            if paragraph.strip():
                story.append(Paragraph(paragraph, content_style))
                story.append(Spacer(1, 0.2 * inch))

        # Generate PDF
        doc.build(story)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_docx(title, content):
        """Generate DOCX document"""
        doc = DocxDocument()

        # Add title
        title_paragraph = doc.add_heading(title, level=1)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Style title
        title_run = title_paragraph.runs[0]
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0, 58, 128)  # KnowHow blue

        doc.add_paragraph()  # Add space

        # Add content
        for paragraph in content.split('\n'):
            if paragraph.strip():
                p = doc.add_paragraph(paragraph)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                # Style content
                for run in p.runs:
                    run.font.size = Pt(12)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


# ============================================================================
# 4. views.py - API Views
# ============================================================================

from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from django.http import FileResponse, HttpResponse
from django.core.files.base import ContentFile
from .models import Document
from .serializers import DocumentSerializer, GenerateDocumentSerializer
from .document_generator import DocumentGenerator

class DocumentViewSet(viewsets.ModelViewSet):
    """ViewSet for Document CRUD operations"""

    permission_classes = [IsAuthenticated]
    serializer_class = DocumentSerializer

    def get_queryset(self):
        """Return documents for current user only"""
        return Document.objects.filter(user=self.request.user)

    def list(self, request):
        """GET /api/documents/ - List all documents"""
        documents = self.get_queryset()
        serializer = self.get_serializer(documents, many=True)
        return Response(serializer.data)

    def retrieve(self, request, pk=None):
        """GET /api/documents/{id}/ - Get single document"""
        try:
            document = self.get_queryset().get(pk=pk)
            serializer = self.get_serializer(document)
            return Response(serializer.data)
        except Document.DoesNotExist:
            return Response(
                {'error': 'Document not found'},
                status=status.HTTP_404_NOT_FOUND
            )

    def destroy(self, request, pk=None):
        """DELETE /api/documents/{id}/ - Delete document"""
        try:
            document = self.get_queryset().get(pk=pk)
            document.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except Document.DoesNotExist:
            return Response(
                {'error': 'Document not found'},
                status=status.HTTP_404_NOT_FOUND
            )

    @action(detail=False, methods=['post'])
    def generate(self, request):
        """POST /api/documents/generate/ - Generate new document"""
        serializer = GenerateDocumentSerializer(data=request.data)

        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

        data = serializer.validated_data

        try:
            # Generate document file
            if data['format'] == 'pdf':
                file_buffer = DocumentGenerator.generate_pdf(
                    data['title'],
                    data['content']
                )
                filename = f"{data['title']}.pdf"
                content_type = 'application/pdf'
            else:  # docx
                file_buffer = DocumentGenerator.generate_docx(
                    data['title'],
                    data['content']
                )
                filename = f"{data['title']}.docx"
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

            # Create document record
            document = Document.objects.create(
                user=request.user,
                session_id=data.get('session_id'),
                title=data['title'],
                content=data['content'],
                format=data['format'],
                file_size=file_buffer.getbuffer().nbytes
            )

            # Save file
            document.file.save(
                filename,
                ContentFile(file_buffer.getvalue()),
                save=True
            )

            # Return response
            return Response({
                'document_id': str(document.id),
                'file_url': document.file_url,
                'title': document.title,
                'format': document.format,
                'created_at': document.created_at.isoformat()
            }, status=status.HTTP_201_CREATED)

        except Exception as e:
            return Response(
                {'error': f'Failed to generate document: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )

    @action(detail=True, methods=['get'])
    def download(self, request, pk=None):
        """GET /api/documents/{id}/download/ - Download document"""
        try:
            document = self.get_queryset().get(pk=pk)

            if not document.file:
                return Response(
                    {'error': 'File not found'},
                    status=status.HTTP_404_NOT_FOUND
                )

            # Set content type based on format
            content_type = (
                'application/pdf' if document.format == 'pdf'
                else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

            # Return file
            response = FileResponse(
                document.file.open('rb'),
                content_type=content_type
            )
            response['Content-Disposition'] = f'attachment; filename="{document.title}.{document.format}"'
            return response

        except Document.DoesNotExist:
            return Response(
                {'error': 'Document not found'},
                status=status.HTTP_404_NOT_FOUND
            )


# ============================================================================
# 5. urls.py - URL Configuration
# ============================================================================

from django.urls import path, include
from rest_framework.routers import DefaultRouter
from .views import DocumentViewSet

router = DefaultRouter()
router.register(r'documents', DocumentViewSet, basename='document')

urlpatterns = [
    path('api/', include(router.urls)),
]


# ============================================================================
# 6. ai_integration.py - AI Document Detection
# ============================================================================

import re
import json

class AIDocumentDetector:
    """Detect when AI should generate documents and format responses"""

    # Georgian keywords for document generation
    DOCUMENT_KEYWORDS_GE = [
        'შექმენი დოკუმენტი',
        'დააგენერირე',
        'გააკეთე დოკუმენტი',
        'მოამზადე',
        'ჩამოტვირთვადი',
        'PDF',
        'DOCX',
        'ხელშეკრულება',
        'განცხადება',
        'საჩივარი',
    ]

    @staticmethod
    def should_generate_document(user_message):
        """Check if user is requesting document generation"""
        message_lower = user_message.lower()
        return any(keyword.lower() in message_lower
                   for keyword in AIDocumentDetector.DOCUMENT_KEYWORDS_GE)

    @staticmethod
    def create_document_marker(title, content, format='pdf'):
        """Create document generation marker for frontend"""
        doc_data = {
            'title': title,
            'content': content,
            'format': format
        }
        return f'<<<GENERATE_DOCUMENT:{json.dumps(doc_data, ensure_ascii=False)}>>>'

    @staticmethod
    def process_ai_response(user_message, ai_response):
        """Process AI response and add document marker if needed"""

        if not AIDocumentDetector.should_generate_document(user_message):
            return ai_response

        # Extract document title from AI response or use default
        title = "Legal_Document"
        if "ხელშეკრულება" in ai_response:
            title = "Service_Agreement"
        elif "განცხადება" in ai_response:
            title = "Application"
        elif "საჩივარი" in ai_response:
            title = "Complaint"

        # Add document marker to response
        marker = AIDocumentDetector.create_document_marker(
            title=title,
            content=ai_response,
            format='pdf'
        )

        # Return response with marker
        return f"{ai_response}\n\n{marker}\n\nდოკუმენტი მზადაა გადმოსაწერად!"


# ============================================================================
# 7. Example Chat View Integration
# ============================================================================

from .ai_integration import AIDocumentDetector

class ChatView(APIView):
    """Chat endpoint with document generation support"""

    permission_classes = [IsAuthenticated]

    def post(self, request):
        user_message = request.data.get('message')
        session_id = request.data.get('sessionId')

        # Get AI response (your existing AI logic)
        ai_response = your_ai_function(user_message, session_id)

        # Process response for document generation
        processed_response = AIDocumentDetector.process_ai_response(
            user_message,
            ai_response
        )

        return Response({
            'content': processed_response,
            'session_id': session_id
        })


# ============================================================================
# 8. requirements.txt - Required Packages
# ============================================================================

# Add these to your requirements.txt:
# reportlab==4.0.7
# python-docx==1.1.0
# djangorestframework==3.14.0
# Pillow==10.1.0


# ============================================================================
# 9. settings.py - Configuration
# ============================================================================

# Add to Django settings.py:

INSTALLED_APPS = [
    # ... other apps
    'rest_framework',
    'corsheaders',
    'your_app',  # Your app with documents
]

MIDDLEWARE = [
    'corsheaders.middleware.CorsMiddleware',
    # ... other middleware
]

# CORS Configuration
CORS_ALLOWED_ORIGINS = [
    "https://knowhow.ge",
    "http://localhost:3000",
]
CORS_ALLOW_CREDENTIALS = True

# CSRF Configuration
CSRF_TRUSTED_ORIGINS = [
    "https://knowhow.ge",
    "http://localhost:3000",
]

# Session Configuration
SESSION_COOKIE_DOMAIN = '.knowhow.ge'
CSRF_COOKIE_DOMAIN = '.knowhow.ge'
SESSION_COOKIE_SECURE = True  # For HTTPS
CSRF_COOKIE_SECURE = True     # For HTTPS

# Media Files (for document storage)
MEDIA_URL = '/media/'
MEDIA_ROOT = os.path.join(BASE_DIR, 'media')


# ============================================================================
# DEPLOYMENT INSTRUCTIONS
# ============================================================================

"""
1. Install requirements:
   pip install reportlab python-docx djangorestframework Pillow

2. Create migrations:
   python manage.py makemigrations
   python manage.py migrate

3. Add URLs to main urls.py:
   from your_app.urls import urlpatterns as doc_urls
   urlpatterns += doc_urls

4. Test endpoints:
   python manage.py runserver

5. Test with frontend:
   - Start frontend: npm start
   - Login to app
   - Try chatting with document request
   - Check if document generates

6. Deploy:
   - Ensure MEDIA_ROOT is configured
   - Set up nginx to serve media files
   - Enable HTTPS for secure cookies
   - Test end-to-end
"""
